from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Bao_cao_cap_nhat_FreshFood_AI_20260619.docx"

SOURCES = [
    ("PHẦN 1. PHÂN TÍCH KHÁC BIỆT", ROOT / "06_phan_tich_khac_biet_bao_cao_va_source.md"),
    ("PHẦN 2. DANH SÁCH NỘI DUNG ĐÃ CẬP NHẬT", ROOT / "07_danh_sach_noi_dung_da_cap_nhat.md"),
    ("PHẦN 3. BÁO CÁO HOÀN CHỈNH", ROOT / "08_bao_cao_hoan_chinh_cap_nhat.md"),
]


def flush_table(document: Document, table_lines: list[str]) -> None:
    if not table_lines:
        return

    rows = []
    for raw in table_lines:
        stripped = raw.strip()
        if not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if not cells:
            continue
        if set("".join(cells).replace("-", "").replace(":", "").strip()) == set():
            continue
        rows.append(cells)

    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            table.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""


def add_markdown(document: Document, text: str) -> None:
    table_buffer: list[str] = []

    def flush_if_needed() -> None:
        nonlocal table_buffer
        flush_table(document, table_buffer)
        table_buffer = []

    for line in text.splitlines():
        stripped = line.rstrip()

        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            continue

        flush_if_needed()

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
            continue

        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.startswith("> "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(stripped[2:].strip())
            run.italic = True
            continue

        paragraph = document.add_paragraph()
        paragraph.add_run(stripped)

    flush_if_needed()


def build() -> Path:
    document = Document()
    document.add_heading("BÁO CÁO CẬP NHẬT DỰ ÁN FRESHFOOD AI", level=0)
    document.add_paragraph("Ngày tạo: 19/06/2026")
    document.add_paragraph(
        "Tài liệu này được sinh từ source code hiện tại và các báo cáo Markdown cập nhật trong thư mục docs."
    )

    for title, path in SOURCES:
        document.add_page_break()
        document.add_heading(title, level=1)
        add_markdown(document, path.read_text(encoding="utf-8"))

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build()
    print(output)
